import os
from docx import Document
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
import requests
import zipfile

# === 1. Funções para converter e extrair notícias ===
def extrair_noticias_heading1(caminho_docx):
    doc = Document(caminho_docx)
    noticias = {}
    atual = []
    contador = 0
    em_noticia = False

    for p in doc.paragraphs:
        estilo = p.style.name.lower()
        texto = p.text.strip()
        if not texto:
            continue

        if estilo == "heading 1":
            if atual:
                contador += 1
                noticias[f"noticia{contador}"] = "\n".join(atual).strip()
                atual = []
            em_noticia = True
            atual.append(texto)
        elif em_noticia:
            atual.append(texto)

    if atual:
        contador += 1
        noticias[f"noticia{contador}"] = "\n".join(atual).strip()

    print(f"Total de notícias encontradas: {len(noticias)}")
    return noticias

def processar_arquivo(caminho_docx):
    if not caminho_docx.lower().endswith('.docx'):
        raise ValueError("O arquivo precisa ser .docx")
    
    print("Extraindo notícias baseadas em Heading 1...")
    noticias = extrair_noticias_heading1(caminho_docx)
    
    return noticias


# === 2. Inicializar ambiente e chain de resumo ===
load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")

prompt = ChatPromptTemplate.from_template(
    "Resuma a notícia: {noticia} em até 100 palavras (não ultrapasse 100 palavras)."
)
chain = prompt | ChatOpenAI() | StrOutputParser()

# === 3. Salvar resumos em arquivo .docx ===
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import requests

# 🔍 Função de busca no Google
def buscar_link_google(titulo, veiculo):
    api_key = os.getenv("GOOGLE_SEARCH_API_KEY")
    cx = os.getenv("GOOGLE_CX")
    
    # Inclui tanto título quanto veículo para maior precisão (sem aspas)
    query = f'"{titulo}"' #{veiculo}'
    
    url = f"https://www.googleapis.com/customsearch/v1?key={api_key}&cx={cx}&q={query}&dateRestrict=d30"

    dominios_bloqueados = [
        "instagram.com", "facebook.com", "twitter.com", "x.com", "linkedin.com"
    ]

    try:
        res = requests.get(url)
        data = res.json()

        total = data.get("searchInformation", {}).get("totalResults")
        print(f"[DEBUG] totalResults: {total}")

        if "items" in data and data["items"]:
            '''primeiro_link = data["items"][0]["link"]
            print(f"[DEBUG] Link retornado: {primeiro_link}")
            return primeiro_link'''
            for item in data["items"]:
                link = item.get("link", "")
                if not any(dom in link for dom in dominios_bloqueados):
                    print(f"[DEBUG] Link válido encontrado: {link}")
                    return link
            print("[DEBUG] Todos os links encontrados foram de redes sociais.")
        else:
            print("[DEBUG] Nenhum item encontrado na resposta.")
    except Exception as e:
        print(f"[ERRO] Falha ao buscar no Google: {e}")
    
    return None

# 🔗 Função para inserir link clicável
def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Estiliza o link (azul e sublinhado)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), "0000FF")
    rPr.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), "single")
    rPr.append(underline)

    new_run.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

# 📝 Função principal exportando título + link + resumo
from docx.shared import Pt, RGBColor

def exportar_resumos_para_word(noticias_dict, resumos_dict, caminho_saida='resumos.docx'):
    doc = Document()
    doc.add_heading('Resumos das Notícias', level=1)

    for i in range(1, len(noticias_dict) + 1):
        noticia_key = f'noticia{i}'
        resumo_key = f'resumo{i}'

        noticia = noticias_dict.get(noticia_key, '')
        resumo = resumos_dict.get(resumo_key, '[Resumo não disponível]')

        linhas = noticia.split('\n')
        titulo = linhas[0] if len(linhas) > 0 else '[Título não encontrado]'
        veiculo = linhas[2] if len(linhas) > 2 else '[Veículo não identificado]'

        link = buscar_link_google(titulo, veiculo)

        # Título com hyperlink (texto clicável e em negrito)
        paragrafo_titulo = doc.add_paragraph()
        if link and link.startswith("http"):
            run = paragrafo_titulo.add_run(titulo)
            run.bold = True
            run.font.size = Pt(11)
            # Faz o título ser clicável
            add_hyperlink(paragrafo_titulo, titulo, link)
            # Remove o run de texto plano (visualmente fica duplicado se não limpar)
            paragrafo_titulo._element.clear_content()
            # Reinsere só o link clicável com o mesmo texto
            add_hyperlink(paragrafo_titulo, titulo, link)
        else:
            run = paragrafo_titulo.add_run(titulo)
            run.bold = True
            run.font.size = Pt(11)

        # Resumo
        p = doc.add_paragraph(resumo)
        p.style.font.size = Pt(11)

        # Aviso se for do Valor Econômico
        if veiculo.startswith("Valor Economico"):
            aviso = doc.add_paragraph("Notícia anexa ao e-mail")
            aviso.style.font.size = Pt(10)
            #aviso.runs[0].italic = True

    doc.save(caminho_saida)
    print(f"\nArquivo Word exportado com sucesso para: {os.path.abspath(caminho_saida)}")

# === Função para identificar notícias do Valor Econômico
def extrair_valor_economico(noticias_dict):
    noticias_valor = {}
    for key, conteudo in noticias_dict.items():
        linhas = conteudo.split('\n')
        if len(linhas) >= 3:
            veiculo = linhas[2].strip()
            if veiculo.startswith("Valor Economico"):
                noticias_valor[key] = {
                    "titulo": linhas[0].strip(),
                    "veiculo": veiculo,
                    "conteudo": '\n'.join(linhas[3:]).strip()
                }
    return noticias_valor

# === Funções para gerar PDFd
from fpdf import FPDF

def criar_pdf(titulo, veiculo, conteudo, caminho):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Fonte padrão Helvetica
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, titulo, ln=True, align="L")

    pdf.ln(4)
    pdf.set_font("Helvetica", "", 11)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 8, f"Fonte: {veiculo}", ln=True)

    pdf.ln(6)
    pdf.set_text_color(0, 0, 0)

    # Quebrar o texto do conteúdo
    pdf.set_font("Helvetica", "", 11)
    for linha in conteudo.split("\n"):
        pdf.multi_cell(0, 6, linha.strip(), align="J")
        pdf.ln(1)

    pdf.output(caminho)

def salvar_noticias_valor_pdf(noticias_valor):
    arquivos = []
    pasta_saida = "pdfs_valor"
    os.makedirs(pasta_saida, exist_ok=True)

    for i, (key, info) in enumerate(noticias_valor.items(), 1):
        nome = f"{i:02d}_{info['titulo'][:40].replace(' ', '_')}.pdf"
        caminho = os.path.join(pasta_saida, nome)
        criar_pdf(info['titulo'], info['veiculo'], info['conteudo'], caminho)
        arquivos.append(caminho)

    return arquivos

# === Função para zip dos PDFs (se necessário)
def compactar_em_zip(lista_caminhos, caminho_zip):
    with zipfile.ZipFile(caminho_zip, 'w') as zipf:
        for caminho in lista_caminhos:
            zipf.write(caminho, os.path.basename(caminho))
    return caminho_zip

# === 4. Executar todo o processo: leitura, resumo e exibição ===
def resumir_noticias(noticias_dict):
    resumos = {}
    for chave, noticia in noticias_dict.items():
        try:
            resumo = chain.invoke({'noticia': noticia})
            resumos[chave.replace("noticia", "resumo")] = resumo
        except Exception as e:
            resumos[chave.replace("noticia", "resumo")] = f"[ERRO] {e}"
    return resumos
